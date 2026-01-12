from flask import Flask, render_template, request, flash, redirect, url_for, send_from_directory
from docx import Document
from textblob import TextBlob
import openai
import os
import json

app = Flask(__name__, static_url_path='/static', static_folder='static')

app.secret_key = '7421'
api_key = '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'

openai.api_key = api_key

# Function to transcribe audio
def transcribe_audio(audio_file):
    # Save the audio file to a temporary location
    temp_file_path = 'temp_audio_file.wav'  # Change the file extension if needed
    audio_file.save(temp_file_path)

    # Open the audio file as a file object
    with open(temp_file_path, 'rb') as file:
        # Transcribe the audio from the file object
        transcription = openai.Audio.transcribe("whisper-1", file)

    # Remove the temporary file after transcription
    os.remove(temp_file_path)

    return transcription['text']


# Function for sentiment analysis
def sentiment_analysis(transcription_text):
    blob = TextBlob(transcription_text)
    polarity = blob.sentiment.polarity
    if polarity > 0:
        return "Positive"
    elif polarity < 0:
        return "Negative"
    else:
        return "Neutral"

# Function to extract key points using GPT-3.5-turbo
def key_points_extraction(transcription_text):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo-16k-0613",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription_text
            }
        ]
    )
    return response['choices'][0]['message']['content']

# Function to extract action items using GPT-3.5-turbo
def action_item_extraction(transcription_text):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo-16k-0613",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in identifying action items. Please review the following text and identify any specific tasks or action items that were assigned or discussed during the meeting."
            },
            {
                "role": "user",
                "content": transcription_text
            }
        ]
    )
    return response['choices'][0]['message']['content']

# Function to save meeting minutes as DOCX
def save_as_docx(transcription_text, minutes, filename):
    doc = Document()

    # Add Meeting Minutes section
    doc.add_heading('Meeting Minutes', level=1)
    doc.add_paragraph(transcription_text)
    doc.add_page_break()

    # Add Abstract Summary section
    doc.add_heading('Abstract Summary', level=1)
    doc.add_paragraph(minutes['abstract_summary'])
    doc.add_page_break()

    # Add Key Points section
    doc.add_heading('Key Points', level=1)
    if minutes['key_points']:
        doc.add_paragraph(minutes['key_points'])
    else:
        doc.add_paragraph("No information available.")
    doc.add_page_break()

    # Add Action Items section
    doc.add_heading('Action Items', level=1)
    if minutes['action_items']:
        doc.add_paragraph(minutes['action_items'])
    else:
        doc.add_paragraph("No information available.")
    doc.add_page_break()

    # Add Sentiment section
    doc.add_heading('Sentiment', level=1)
    doc.add_paragraph(minutes['sentiment'])

    doc.save(filename)

# Function to extract abstract summary using GPT-3.5-turbo
def abstract_summary_extraction(transcription_text):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo-16k-0613",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription_text
            }
        ]
    )
    return response['choices'][0]['message']['content']

# Function to extract meeting minutes
def meeting_minutes(transcription_text):
    abstract_summary = abstract_summary_extraction(transcription_text)
    key_points = key_points_extraction(transcription_text)
    action_items = action_item_extraction(transcription_text)
    sentiment = sentiment_analysis(transcription_text)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

# Function to handle pasted text
def handle_pasted_text(text):
    return text

@app.route('/', methods=['GET', 'POST'])
def index():
    minutes = None

    if request.method == 'POST':
        input_type = request.form['input_type']

        if input_type == 'audio':
            # Check if the audio file is uploaded
            if 'audio_file' not in request.files:
                flash("Please upload an audio file", "error")
                return redirect(request.url)

            audio_file = request.files['audio_file']

            # Check if the file is empty
            if audio_file.filename == '':
                flash("No selected file", "error")
                return redirect(request.url)

            # Check if the file is an audio file with a supported format
            allowed_formats = {'m4a', 'mp3', 'webm', 'mp4', 'mpga', 'wav', 'mpeg', 'ogg', 'oga', 'flac'}
            if audio_file.filename.split('.')[-1].lower() not in allowed_formats:
                flash("Invalid file format. Supported formats: {}".format(', '.join(allowed_formats)), "error")
                return redirect(request.url)

            # Transcribe the audio
            try:
                transcription_text = transcribe_audio(audio_file)
            except openai.error.OpenAIError:
                flash("Error during audio transcription. Please try again later.", "error")
                return redirect(request.url)

            # Extract meeting minutes
            minutes = meeting_minutes(transcription_text)

            # Convert minutes dictionary to JSON data
            minutes_json = json.dumps({
                'transcription_text': transcription_text,
                'abstract_summary': minutes['abstract_summary'],
                'key_points': minutes['key_points'],
                'action_items': minutes['action_items'],
                'sentiment': minutes['sentiment']
            })

            filename = audio_file.filename.split('.')[0] + '.docx'
        
            save_as_docx(transcription_text, minutes, filename)
            flash("Meeting minutes extracted and saved successfully!", "success")

            # Add filename to render_template to pass it to the 'result.html'
            return render_template('result.html', minutes=minutes, filename=filename)

        elif input_type == 'text':
            pasted_text = request.form['pasted_text']
            transcription_text = handle_pasted_text(pasted_text)

            # Extract meeting minutes
            minutes = meeting_minutes(transcription_text)

            # Convert minutes dictionary to JSON data
            minutes_json = json.dumps({
                'transcription_text': transcription_text,
                'abstract_summary': minutes['abstract_summary'],
                'key_points': minutes['key_points'],
                'action_items': minutes['action_items'],
                'sentiment': minutes['sentiment']
            })

            filename = 'text_based.docx'
        
            save_as_docx(transcription_text, minutes, filename)
            flash("Meeting minutes extracted and saved successfully!", "success")

            # Add filename to render_template to pass it to the 'result.html'
            return render_template('result.html', minutes=minutes, filename=filename)

    return render_template('index.html')

# New route to handle the download of the DOCX file
@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory('.', filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
